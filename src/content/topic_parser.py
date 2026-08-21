from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


# Explicit section/category headings from the current topics.docx format.
# These are used as category context, not as individual post topics.
KNOWN_CATEGORIES = {
    "aws architecture": "aws",
    "java internals": "java",
    "spring internals": "spring",
    "database architecture": "sql",
    "kafka deep dive": "kafka",
    "redis": "redis",
    "docker & kubernetes": "docker",
    "production engineering": "production",
}


# Keyword-based fallback for topics that do not come from a known heading.
CATEGORY_RULES = {
    "kafka": [
        "kafka", "consumer group", "producer", "partition", "offset",
        "rebalance", "isr", "leader election", "schema registry"
    ],
    "redis": [
        "redis", "cache", "caching", "ttl", "eviction", "redlock",
        "sentinel", "redis cluster", "pub/sub", "streams"
    ],
    "spring": [
        "spring boot", "spring ", "transactional", "spring security",
        "spring cloud", "actuator", "bean", "dependency injection",
        "hibernate", "lazy loading", "open session in view"
    ],
    "java": [
        "java", "jvm", "hashmap", "garbage collection", "gc ", "g1 gc",
        "zgc", "escape analysis", "jit", "bytecode", "reflection",
        "java nio", "virtual thread", "project loom", "record",
        "sealed class", "pattern matching", "stream api",
        "parallel stream", "completablefuture", "thread dump", "heap dump"
    ],
    "aws": [
        "aws", "api gateway", "lambda", "ecs", "eks", "sqs", "sns",
        "eventbridge", "rds", "dynamodb", "elasticache", "vpc",
        "route 53", "s3", "ec2", "cloudfront", "iam"
    ],
    "mongodb": [
        "mongodb", "mongo", "document database", "aggregation"
    ],
    "docker": [
        "docker", "kubernetes", "pod", "deployment", "replicaset",
        "ingress", "configmap", "hpa", "namespace", "node vs pod",
        "cluster architecture", "docker networking", "docker volumes"
    ],
    "microservices": [
        "microservice", "circuit breaker", "saga", "bulkhead",
        "retry pattern", "service discovery", "distributed system"
    ],
    "sql": [
        "sql", "mysql", "database", "b+ tree", "index", "clustered",
        "non clustered", "replication", "partitioning", "connection pool",
        "acid", "cap theorem", "base architecture", "two phase commit",
        "optimistic locking", "pessimistic locking", "read replica",
        "leader follower", "master slave"
    ],
    "security": [
        "jwt", "oauth", "authentication", "authorization", "security",
        "token", "credential"
    ],
    "production": [
        "memory leak", "thread dump analysis", "heap dump analysis",
        "cpu profiling", "gc logs", "high cpu", "database bottleneck",
        "redis failure", "kafka failure", "api timeout", "slow sql",
        "circuit breaker in action", "retry pattern", "bulkhead pattern",
        "observability", "prometheus", "grafana", "elk", "opentelemetry"
    ],
}


def clean_text(value: str) -> str:
    """Normalize whitespace and remove invisible formatting artifacts."""
    value = value.replace("\u00a0", " ")
    return " ".join(value.strip().split())


def normalize_category(value: str) -> str:
    return clean_text(value).casefold()


def normalize_topic(value: str) -> str:
    """
    Convert bullet/numbered lines into the actual topic text.

    Examples:
      '• API Gateway' -> 'API Gateway'
      '1. API Gateway' -> 'API Gateway'
      '1) API Gateway' -> 'API Gateway'
      '- API Gateway' -> 'API Gateway'
    """
    text = clean_text(value)

    # Remove common bullet characters.
    text = re.sub(r"^[•●▪◦‣\-–—]\s*", "", text)

    # Remove simple numbered list prefixes.
    text = re.sub(r"^\d+\s*[\.\)]\s*", "", text)

    return clean_text(text)


def is_section_heading(text: str) -> bool:
    """
    Section headings in the supplied topics.docx look like:
      1. AWS Architecture
      2. Java Internals
      ...
    """
    normalized = clean_text(text).casefold()

    # Exact known headings.
    if normalized in KNOWN_CATEGORIES:
        return True

    # Numbered heading pattern such as:
    # 1. AWS Architecture
    numbered = re.match(r"^\d+\s*[\.\)]\s*(.+)$", normalized)
    if numbered:
        heading = clean_text(numbered.group(1))
        if heading in KNOWN_CATEGORIES:
            return True

    return False


def extract_heading_category(text: str) -> str | None:
    normalized = clean_text(text)
    key = normalized.casefold()

    if key in KNOWN_CATEGORIES:
        return KNOWN_CATEGORIES[key]

    match = re.match(r"^\d+\s*[\.\)]\s*(.+)$", normalized)
    if match:
        heading = clean_text(match.group(1)).casefold()
        return KNOWN_CATEGORIES.get(heading)

    return None


def detect_category(topic: str, section_category: str | None = None) -> str:
    """
    Prefer the category supplied by the Word document section.
    Otherwise fall back to keyword detection.
    """
    if section_category:
        return section_category

    text = topic.casefold()

    # Production is more specific than generic Java/Kafka/Redis mentions.
    for category in ("production", "microservices", "security", "kafka", "redis", "spring", "java", "aws", "docker", "mongodb", "sql"):
        for keyword in CATEGORY_RULES[category]:
            if keyword in text:
                return category

    return "generic"


def split_multi(value: str) -> list[str]:
    """For optional rich Word-table fields separated by semicolons."""
    if not value:
        return []
    return [clean_text(part) for part in value.split(";") if clean_text(part)]


def _base_item(topic: str, section_category: str | None) -> dict[str, Any]:
    category = detect_category(topic, section_category)
    return {
        "topic": topic,
        "category": category,
        "section_category": section_category,
        "overview": "",
        "key_concepts": [],
        "diagram": "",
        "scenarios": [],
        "best_practices": [],
        "use_cases": [],
    }


def _read_paragraph_topics(doc: Document) -> list[dict[str, Any]]:
    """
    Parse the current topics.docx paragraph format:

        1. AWS Architecture
        • API Gateway
        • Lambda
        ...
        2. Java Internals
        • JVM Startup Flow
        ...

    Section headings are ignored as posts but retained as category context.
    """
    items: list[dict[str, Any]] = []
    current_section: str | None = None

    for paragraph in doc.paragraphs:
        raw = clean_text(paragraph.text)
        if not raw:
            continue

        heading_category = extract_heading_category(raw)
        if heading_category:
            current_section = heading_category
            continue

        topic = normalize_topic(raw)
        if not topic:
            continue

        # Ignore accidental plain-text lines that are numbered headings
        # even when the heading is not in our known-category map.
        if re.match(r"^\d+\s*[\.\)]\s+", raw):
            continue

        items.append(_base_item(topic, current_section))

    return items


def _read_rich_tables(doc: Document) -> list[dict[str, Any]]:
    """
    Optional table format:

    Topic | Category | Overview | Key Concepts | Diagram | Scenarios |
    Best Practices | Use Cases

    Semicolon separates multiple values inside a cell.
    """
    items: list[dict[str, Any]] = []

    for table in doc.tables:
        if not table.rows:
            continue

        headers = [normalize_category(cell.text) for cell in table.rows[0].cells]
        if "topic" not in headers:
            continue

        index = {header: i for i, header in enumerate(headers)}
        if "topic" not in index:
            continue

        for row in table.rows[1:]:
            cells = [clean_text(cell.text) for cell in row.cells]

            def get(name: str) -> str:
                idx = index.get(name)
                if idx is None or idx >= len(cells):
                    return ""
                return cells[idx]

            topic = normalize_topic(get("topic"))
            if not topic:
                continue

            category_value = normalize_category(get("category"))
            item = _base_item(topic, category_value or None)
            item["overview"] = get("overview")
            item["key_concepts"] = split_multi(get("key concepts"))
            item["diagram"] = get("diagram")
            item["scenarios"] = split_multi(get("scenarios"))
            item["best_practices"] = split_multi(get("best practices"))
            item["use_cases"] = split_multi(get("use cases"))
            items.append(item)

    return items


def _deduplicate(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Preserve document order and remove duplicate topics case-insensitively."""
    seen: set[str] = set()
    unique: list[dict[str, Any]] = []

    for item in items:
        key = normalize_topic(item["topic"]).casefold()
        if not key or key in seen:
            continue

        seen.add(key)
        unique.append(item)

    return unique


def read_topics(path: str) -> list[dict[str, Any]]:
    """
    Read topics.docx.

    Supported formats:
      1) Current grouped bullet-list Word file.
      2) Optional rich table format.

    The grouped-list format is intentionally the primary source.
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"Topics file not found: {file_path}")

    doc = Document(file_path)

    paragraph_items = _read_paragraph_topics(doc)
    table_items = _read_rich_tables(doc)

    # Prefer paragraphs when the file uses the current format.
    # Rich tables are additive for users who choose the advanced format.
    items = paragraph_items + table_items

    return _deduplicate(items)
